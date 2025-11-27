import os
import re
from collections import Counter
from docx import Document
import pdfplumber
from RankCvsSemantic import rank_cvs_semantic
from RankCvsSemanticWeighted import rank_cvs_semantic_weighted
import sys
from Translate import translate_text

doc = Document()

# ---- Funktion för att läsa text från PDF ----
def read_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

# ---- Funktion för att läsa text från DOCX ----
def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# ---- Extrahera text från CV-filer ----
def extract_cv_text(folder_path):
    cv_texts = {}
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.lower().endswith(".pdf"):
            cv_texts[filename] = read_pdf(file_path)
        elif filename.lower().endswith(".docx"):
            cv_texts[filename] = read_docx(file_path)
    return cv_texts

# ---- Beräkna matchningspoäng ----
def calculate_score(cv_text, job_description):
    # Gör allt lowercase
    cv_text = cv_text.lower()
    job_description = job_description.lower()
    
    # Extrahera ord
    job_words = re.findall(r"\w+", job_description)
    cv_words = re.findall(r"\w+", cv_text)
    
    # Räkna förekomster
    cv_counter = Counter(cv_words)
    
    score = sum(cv_counter[word] for word in job_words)
    return score

# ---- Rangordna CV:n ----
def rank_cvs(folder_path, job_description):
    cv_texts = extract_cv_text(folder_path)
    scores = {filename: calculate_score(text, job_description) for filename, text in cv_texts.items()}
    ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    return ranked


def read_weighted_keywords(file_path):
    weighted_keywords = {}
    with open(file_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line and "=" in line:
                key, value = line.split("=", 1)
                try:
                    weighted_keywords[key.strip()] = float(value.strip())
                except ValueError:
                    print(f"Ogiltigt värde för {key}: {value}")
                    doc.add_paragraph(f"Ogiltigt värde för {key}: {value}")
    return weighted_keywords

def read_textfile(filepath: str) -> str:
    """
    Läser in en textfil och returnerar innehållet som en sträng.
    
    :param filväg: Sökvägen till textfilen (t.ex. 'C:/Dev/test.txt')
    :return: Filens textinnehåll som en sträng
    """
    try:
        with open(filepath, "r", encoding="utf-8") as fil:
            innehåll = fil.read()
        return innehåll
    except FileNotFoundError:
        print(f"❌ Filen '{filepath}' hittades inte.")
        doc.add_paragraph(f"❌ Filen '{filepath}' hittades inte.")
        return ""
    except Exception as e:
        print(f"⚠️ Ett fel uppstod vid läsning av filen: {e}")
        doc.add_paragraph(f"⚠️ Ett fel uppstod vid läsning av filen: {e}")
        return ""

# ---- Exempelanrop ----
if __name__ == "__main__":


    if len(sys.argv) < 5:
        szTmp = "Usage: python RangordnaCV.py <cv_folder> <job_description> <keywords_file>"
        doc.add_paragraph(szTmp)
        cv_folder = r"C:\Dev\TestData\CV"
        job_description = r"C:\Dev\TestData\jobdesc.txt"
        keywords_file = r"C:\Dev\TestData\weightedkeywords.txt"
        working_lnguage = "no"
        log_file = "RangordnaCV_log.docx"
    else:
        cv_folder = sys.argv[1]
        job_description = sys.argv[2]
        keywords_file = sys.argv[3]
        working_lnguage = sys.argv[4]

    if len(sys.argv) > 5:
        log_file = sys.argv[5]


    doc.add_paragraph(f"CV-mapp: {cv_folder}")
    doc.add_paragraph(f"Jobbeskrivning: {job_description}")
    doc.add_paragraph(f"Nyckelordsfil: {keywords_file}")
    folder = cv_folder # "C:\Dev\TestData\CV" # "cv_mapp"  # Ange din mapp med CV:n
    #job_desc = """Vi söker en senior Python-utvecklare med erfarenhet av molnlösningar och AI."""
    #job_desc = """Vi söker en mekanist med erfarenhet av special stål och plast."""
    job_desc = read_textfile(job_description) #"C:\Dev\TestData\jobdesc.txt")  # Läs jobbeskrivningen från en textfil

    translateDoc = working_lnguage != "no"

    source_lang='auto'
    target_lang='sv' #default value
    if working_lnguage == "en":
        target_lang='en'
    elif working_lnguage == "sv":
        target_lang='sv'

    if translateDoc:
        job_desc = translate_text(job_desc, source_lang, target_lang)
    if not job_desc:
        szTmp = "Ingen jobbeskrivning angiven. Avslutar."
        doc.add_paragraph(szTmp)
    else:
        szTmp = f"Filen innehåller: {job_desc}"
        doc.add_paragraph(szTmp)
        #ranking = rank_cvs(folder, job_desc)
        #ranking =  rank_cvs_semantic(folder, job_desc)
        #Exempel: weightedKeywords={"python":2,"molnlösningar":3,"ai":4}, alpha=0.7, beta=0.3
        myWeightedKeywords=read_weighted_keywords(keywords_file) #"C:\Dev\TestData\weightedkeywords.txt")
        doc.add_paragraph(f"Viktade nyckelord: {myWeightedKeywords}")   

        ranking = rank_cvs_semantic_weighted(folder, job_desc, myWeightedKeywords, translateDoc, source_lang, target_lang)
        print("Rangordning av CV:n:")
        doc.add_paragraph("Rangordning av CV:n:")
        for filename, score in ranking:
            print(f"{filename}: {score}")
            szTmp = f"{filename}: {score}"
            doc.add_paragraph(szTmp)

        doc.save(log_file)
